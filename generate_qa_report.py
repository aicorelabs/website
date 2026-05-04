from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# Add Title
title = doc.add_heading('NeuroChecklist Staging QA Report', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('This document serves as a comprehensive bug-tracking and QA audit report generated for the NeuroChecklist staging environment. It details front-end visual issues, severe code-loop bugs, and heuristic testing results of the AI Chatbot.\n')

doc.add_heading('Bug 1: PostHog Analytics Infinite Retry Loop (Critical)', 1)
p = doc.add_paragraph('Severity: ')
run = p.add_run('High / Blocking')
run.bold = True
doc.add_paragraph('Location: All Pages (Global Header/Footer script)\n')
doc.add_paragraph('Description: When a user accesses the site with standard privacy extensions or ad-blockers (which natively block "eu.i.posthog.com" analytics traffic), the site\'s internal logic enters an infinite retry loop. The console is immediately flooded with hundreds of identical "ERR_BLOCKED_BY_CLIENT" errors. Without exponential backoff built into the analytics integration, this loop drains system memory and rapidly drains user battery life, eventually causing browser crashes on lower-end devices.')
try:
    doc.add_picture(r'C:\Users\TARIBOTHEGLORYPRIEST\.gemini\antigravity\brain\fed5c8fd-368d-4d34-bf2e-a12d8c4115ae\media__1775671293507.png', width=Inches(6.0))
except:
    pass

doc.add_heading('Bug 2: Massive Mobile Horizontal Overflow', 1)
p = doc.add_paragraph('Severity: ')
run = p.add_run('High')
run.bold = True
doc.add_paragraph('Location: Global Mobile Viewport (< 600px width)\n')
doc.add_paragraph('Description: The website layout completely breaks boundaries on mobile devices. A rigid fixed-width element on the main page pushes the viewport width well past 100vw, creating a massive, blank white gutter on the right side of the screen (occupying ~25% of the visible area). This forces users to horizontally scroll, breaking standard mobile UX conventions. Additionally, the top search bar is cropped off-center as a result of this break.')
try:
    doc.add_picture(r'C:\Users\TARIBOTHEGLORYPRIEST\.gemini\antigravity\brain\fed5c8fd-368d-4d34-bf2e-a12d8c4115ae\media__1775672796785.png', width=Inches(6.0))
except:
    pass

doc.add_heading('Bug 3: Scalability Clamp on AI Chatbot UI (Mobile)', 1)
p = doc.add_paragraph('Severity: ')
run = p.add_run('Medium')
run.bold = True
doc.add_paragraph('Location: Chatbot Mock UI Layout on Mobile\n')
doc.add_paragraph('Description: On mobile sizes, the inner text of the "Neurochecklists AI Research Assistant" preview cards (e.g., "Based on McDonald 2017 criteria...") becomes completely crushed against the rounded card borders. The text size does not scale proportionately to its container shrinkage, creating unreadable, overlapping visual bounds.')

doc.add_heading('Bug 4: Missing Dashboard UI Icons', 1)
p = doc.add_paragraph('Severity: ')
run = p.add_run('Low / Visual Bug')
run.bold = True
doc.add_paragraph('Location: Axon AI Chat Interface (Top Right Header)\n')
doc.add_paragraph('Description: In the main AI chat window, there are three empty square border outlines natively rendered in the top-right corner. It appears the SVG icons meant to exist within these buttons (likely Share, Export, or Options) failed to load or were completely omitted from the shipped staging code.')
try:
    doc.add_picture(r'C:\Users\TARIBOTHEGLORYPRIEST\.gemini\antigravity\brain\fed5c8fd-368d-4d34-bf2e-a12d8c4115ae\media__1775673639185.png', width=Inches(6.0))
except:
    pass

doc.add_heading('Bug 5: Critical Search Bar Hijack (Enter Key Error)', 1)
p = doc.add_paragraph('Severity: ')
run = p.add_run('Critical / Blocking')
run.bold = True
doc.add_paragraph('Location: Global Header Search Bar\n')
doc.add_paragraph('Description: There is a catastrophic event-listener bug tied to the global search bar. When a user types a query (e.g., "Alzheimer") and presses the Enter key on their keyboard, the search fails to execute. Instead, the Enter key globally triggers the site\'s "Upgrade/Pricing" gateway modal. Furthermore, the search auto-complete dropdown menu fails to dismiss, remaining permanently stuck overlaid on top of the pricing modal. This completely breaks the primary search-and-discovery UX funnel.')
try:
    doc.add_picture(r'C:\Users\TARIBOTHEGLORYPRIEST\.gemini\antigravity\brain\fed5c8fd-368d-4d34-bf2e-a12d8c4115ae\media__1775746142171.png', width=Inches(6.0))
except:
    pass

doc.add_heading('Chatbot Capability Summary (PASS)', 1)
doc.add_paragraph('Status: PASSED\n')
doc.add_paragraph("Validation: The Axon AI assistant was successfully stress-tested using both domain-specific queries and out-of-scope garbage inputs (e.g. asking for capital cities). The AI guardrails performed flawlessly. It refused to answer non-medical queries and politely guided the user back to the primary neurology workflow, exactly as designed.")

# Saving directly to the Zeffron folder where they moved it last time
output_path = r'c:\Users\TARIBOTHEGLORYPRIEST\Desktop\Zeffron\Zeffron Website_fresh\NeuroChecklist_QA_Audit_Report.docx'
doc.save(output_path)
print(f"Successfully generated updated QA Report at: {output_path}")
